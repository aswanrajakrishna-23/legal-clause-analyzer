from flask import Flask, render_template, request, jsonify, send_file
from analyzer import analyze_clauses
import os
import json
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import io

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/analyze', methods=['POST'])
def analyze():
    data = request.get_json()
    text = data.get('text', '')
    clauses = analyze_clauses(text)
    return jsonify(clauses)

@app.route('/download/pdf', methods=['POST'])
def download_pdf():
    data = request.get_json()
    clauses = data.get('clauses', [])
    
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 50
    
    p.setFont("Helvetica-Bold", 16)
    p.drawString(50, y, "Legal Clause Analysis Report")
    y -= 30
    
    p.setFont("Helvetica", 12)
    for clause in clauses:
        if y < 50:
            p.showPage()
            y = height - 50
            p.setFont("Helvetica", 12)
            
        p.drawString(50, y, f"Type: {clause['type']}")
        y -= 20
        p.drawString(50, y, f"Condition: {clause['condition']}")
        y -= 20
        p.drawString(50, y, f"Consequence: {clause['consequence']}")
        y -= 30
        
    p.save()
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name='analysis_report.pdf', mimetype='application/pdf')

@app.route('/download/docx', methods=['POST'])
def download_docx():
    data = request.get_json()
    clauses = data.get('clauses', [])
    
    document = Document()
    document.add_heading('Legal Clause Analysis Report', 0)
    
    for clause in clauses:
        document.add_heading(clause['type'], level=1)
        document.add_paragraph(f"Condition: {clause['condition']}")
        document.add_paragraph(f"Consequence: {clause['consequence']}")
        
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name='analysis_report.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    app.run(debug=True)
